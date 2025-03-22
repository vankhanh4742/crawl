import os
import re
import time
import uuid
import json
import logging
import random
import threading
import requests
import concurrent.futures
import tempfile

from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options

# Thư viện để tạo file Word
from docx import Document
from docx.shared import Inches

# Thư viện xử lý XML và HTML
from bs4 import BeautifulSoup, NavigableString, Tag, Comment
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Cấu hình logging để ghi lỗi ra file error.log
logging.basicConfig(
    filename="error.log",
    level=logging.ERROR,
    format="%(asctime)s %(levelname)s: %(message)s"
)

# -----------------------
# 1. Hàm làm sạch tên file
def sanitize_filename(filename):
    filename = re.sub(r'[\\/:*?"<>|]', "", filename)
    filename = filename.strip().replace(" ", "_")
    return filename

# -----------------------
# 2. Hàm chuyển đổi MathML thành chuỗi biểu diễn công thức đơn giản
def process_mathml(elem):
    if isinstance(elem, NavigableString):
        return str(elem)
    if not isinstance(elem, Tag):
        return ""
    if elem.name == "math":
        return "".join(process_mathml(child) for child in elem.children)
    elif elem.name == "msqrt":
        return "√(" + "".join(process_mathml(child) for child in elem.children) + ")"
    elif elem.name == "mfrac":
        children = [child for child in elem.children if isinstance(child, Tag)]
        if len(children) >= 2:
            numerator = process_mathml(children[0])
            denominator = process_mathml(children[1])
            return "(" + numerator + ")/(" + denominator + ")"
        else:
            return ""
    elif elem.name == "msup":
        children = [child for child in elem.children if isinstance(child, (Tag, NavigableString))]
        if len(children) >= 2:
            base = process_mathml(children[0])
            sup = process_mathml(children[1])
            return base + "^(" + sup + ")"
        else:
            return ""
    elif elem.name == "msub":
        children = [child for child in elem.children if isinstance(child, (Tag, NavigableString))]
        if len(children) >= 2:
            base = process_mathml(children[0])
            sub = process_mathml(children[1])
            return base + "_(" + sub + ")"
        else:
            return ""
    elif elem.name == "mrow":
        return "".join(process_mathml(child) for child in elem.children)
    elif elem.name in ["mi", "mn", "mo"]:
        return elem.get_text()
    else:
        return "".join(process_mathml(child) for child in elem.children)

# -----------------------
# 3. Hàm chèn hình ảnh inline vào một run
def add_inline_image(run, image_path, img_width, img_height):
    """
    Chèn hình ảnh từ file image_path vào run dưới dạng inline.
    Sử dụng add_picture của python-docx để chèn ảnh theo kích thước (img_width, img_height) tính theo Inches.
    """
    run.add_picture(image_path, width=Inches(img_width / 96), height=Inches(img_height / 96))

# -----------------------
# 4. Hàm lọc bỏ các thẻ không mong muốn (script, style, ins, iframe, noscript) và comment
def clean_soup(soup):
    for tag in soup.find_all(["script", "style", "ins", "iframe", "noscript"]):
        tag.decompose()
    for comment in soup.find_all(text=lambda text: isinstance(text, Comment)):
        comment.extract()
    return soup

# -----------------------
# 5. Hàm chuyển đổi nội dung HTML sang đoạn văn trong file Word
def add_html_to_paragraph(paragraph, element):
    """
    Duyệt qua các phần tử con của element (có thể là <p>, <span>, …)
    và thêm vào đoạn văn (paragraph) của tài liệu Word.
    Xử lý các thẻ: <sup>, <sub>, <strong>, <em>, <br>, <math>, và đặc biệt <img>.
    """
    children = list(element.children)
    for idx, child in enumerate(children):
        if isinstance(child, NavigableString):
            paragraph.add_run(child)
        elif isinstance(child, Tag):
            if child.name == "sup":
                run = paragraph.add_run(child.get_text())
                run.font.superscript = True
            elif child.name == "sub":
                run = paragraph.add_run(child.get_text())
                run.font.subscript = True
            elif child.name == "strong":
                run = paragraph.add_run(child.get_text())
                run.bold = True
            elif child.name == "em":
                run = paragraph.add_run(child.get_text())
                run.italic = True
            elif child.name == "br":
                paragraph.add_run("\n")
            elif child.name == "math":
                math_str = process_mathml(child)
                paragraph.add_run(math_str)
                if idx < len(children) - 1:
                    next_child = children[idx + 1]
                    if isinstance(next_child, Tag) and next_child.name == "math":
                        paragraph.add_run(" ")
            elif child.name == "img":
                # Xử lý chèn hình ảnh inline
                src = child.get("src")
                if src:
                    try:
                        response = requests.get(src, stream=True)
                        if response.status_code == 200:
                            folder_temp = "downloads"
                            if not os.path.exists(folder_temp):
                                os.makedirs(folder_temp)
                            temp_img_path = os.path.join(folder_temp, f"temp_inline_{uuid.uuid4().hex}.jpg")
                            with open(temp_img_path, "wb") as f:
                                for chunk in response.iter_content(chunk_size=8192):
                                    f.write(chunk)
                            try:
                                img_width = int(child.get("width", 200))
                            except:
                                img_width = 200
                            try:
                                img_height = int(child.get("height", 200))
                            except:
                                img_height = 200
                            run = paragraph.add_run()
                            add_inline_image(run, temp_img_path, img_width, img_height)
                            os.remove(temp_img_path)
                        else:
                            paragraph.add_run("[Error loading image]")
                    except Exception as ex:
                        paragraph.add_run("[Exception loading image]")
            else:
                add_html_to_paragraph(paragraph, child)

# -----------------------
# Danh sách user-agent để chọn ngẫu nhiên cho mỗi trình duyệt
USER_AGENTS = [
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/115.0.0.0 Safari/537.36",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 13_2) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/15.2 Safari/605.1.15",
    "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/116.0.0.0 Safari/537.36",
    # Thêm user-agent khác nếu cần
]

# -----------------------
# Hàm tạo driver mới với cấu hình an toàn và user-agent ngẫu nhiên
def create_driver():
    chrome_options = Options()
    chrome_options.add_argument("--headless")
    chrome_options.add_argument("--disable-gpu")
    chrome_options.add_argument("--window-size=1920x1080")
    chrome_options.add_argument("--log-level=3")
    chrome_options.add_argument("--disable-software-rasterizer")
    chrome_options.add_argument("--disable-dev-shm-usage")
    chrome_options.add_argument("--no-sandbox")
    
    user_agent = random.choice(USER_AGENTS)
    chrome_options.add_argument(f"user-agent={user_agent}")
    
    # Sử dụng tempfile.mkdtemp() để tạo thư mục user data duy nhất cho mỗi phiên
    unique_user_data_dir = tempfile.mkdtemp()
    chrome_options.add_argument(f"--user-data-dir={unique_user_data_dir}")
    
    service = Service(r"/usr/local/bin/chromedriver")
    driver = webdriver.Chrome(service=service, options=chrome_options)
    return driver

# -----------------------
# Biến và lock để lưu tiến trình crawl
progress_lock = threading.Lock()
processed_links = set()

# Load tiến trình đã xử lý từ file nếu có
PROGRESS_FILE = "processed_links-toan.txt"
if os.path.exists(PROGRESS_FILE):
    with open(PROGRESS_FILE, "r", encoding="utf-8") as f:
        for line in f:
            processed_links.add(line.strip())

# -----------------------
# Hàm xử lý crawl cho một link cụ thể với cơ chế retry cho việc mở URL
def crawl_link(subject_name, folder_save, href, text):
    with progress_lock:
        if href in processed_links:
            print(f"  >> Link {href} đã được xử lý, bỏ qua.")
            return

    pdf_file = os.path.join(folder_save, sanitize_filename(text) + ".pdf")
    docx_file = os.path.join(folder_save, sanitize_filename(text) + ".docx")
    if os.path.exists(pdf_file) or os.path.exists(docx_file):
        print(f"  >> File cho bài '{text}' đã tồn tại, bỏ qua.")
        with progress_lock:
            processed_links.add(href)
            with open(PROGRESS_FILE, "a", encoding="utf-8") as f:
                f.write(href + "\n")
        return

    max_attempts = 3
    attempt = 0
    driver = None
    while attempt < max_attempts:
        try:
            driver = create_driver()
            driver.get(href)
            break  # Nếu mở URL thành công thì thoát vòng lặp
        except Exception as e:
            attempt += 1
            logging.error(f"Lỗi mở URL {href} (attempt {attempt}): {e}", exc_info=True)
            if driver:
                try:
                    driver.quit()
                except Exception:
                    pass
            time.sleep(random.uniform(2, 4))
    if attempt == max_attempts:
        print(f"  >> Không mở được URL sau {max_attempts} lần: {href}")
        return

    try:
        time.sleep(random.uniform(2, 4))
        driver.execute_script("window.scrollTo(0, document.body.scrollHeight/2);")
        time.sleep(random.uniform(1, 2))

        # Thử tìm nút download PDF
        try:
            download_button = WebDriverWait(driver, 5).until(
                EC.presence_of_element_located((By.ID, "btn-download-md"))
            )
            download_url = download_button.get_attribute("href")
            print(f"  >> Tìm thấy nút download. Link download: {download_url}")
            response = requests.get(download_url, stream=True, headers={"User-Agent": random.choice(USER_AGENTS)})
            if response.status_code == 200:
                if not os.path.exists(folder_save):
                    os.makedirs(folder_save)
                file_name = sanitize_filename(text) + ".pdf"
                file_path = os.path.join(folder_save, file_name)
                with open(file_path, "wb") as file:
                    for chunk in response.iter_content(chunk_size=8192):
                        file.write(chunk)
                print(f"  >> File PDF đã được tải về thành công tại: {file_path}")
            else:
                print("  >> Lỗi khi tải file, status code:", response.status_code)
        except Exception as e:
            print(f"  >> Không tìm thấy nút download tại trang {href}: {e}")
            print("  >> Tiến hành crawl nội dung trang từ thẻ 'div#content-post' và lưu vào file Word...")
            try:
                try:
                    main_content = driver.find_element(By.CSS_SELECTOR, "div#content-post")
                except Exception as ex:
                    print("Không tìm thấy thẻ 'div#content-post'. Sử dụng thẻ body. Lỗi:", ex)
                    main_content = driver.find_element(By.TAG_NAME, "body")
                content_html = main_content.get_attribute("innerHTML")
                soup = BeautifulSoup(content_html, "lxml")
                soup = clean_soup(soup)

                doc = Document()
                doc.add_heading(text, level=1)

                for elem in soup.contents:
                    if isinstance(elem, Tag):
                        p = doc.add_paragraph()
                        add_html_to_paragraph(p, elem)
                    elif isinstance(elem, NavigableString):
                        doc.add_paragraph(elem)

                if not os.path.exists(folder_save):
                    os.makedirs(folder_save)
                file_name = sanitize_filename(text) + ".docx"
                file_path = os.path.join(folder_save, file_name)
                doc.save(file_path)
                print(f"  >> Nội dung trang đã được lưu vào file Word tại: {file_path}")
            except Exception as crawl_e:
                logging.error(f"Lỗi khi crawl nội dung tại {href}: {crawl_e}", exc_info=True)
                print("  >> Lỗi khi crawl nội dung trang, bỏ qua bài này.")
    finally:
        try:
            driver.quit()
        except Exception:
            pass

    # Delay ngẫu nhiên sau khi xử lý mỗi link
    time.sleep(random.uniform(2, 4))
    with progress_lock:
        processed_links.add(href)
        with open(PROGRESS_FILE, "a", encoding="utf-8") as f:
            f.write(href + "\n")

# -----------------------
# Main: Xử lý danh sách subject và crawl các link trong từng subject
def main():
    # Đọc danh sách subjects từ file updated_links.json
    with open("Toan10-link.json", "r", encoding="utf-8") as f:
        subjects = json.load(f)

    # Sử dụng ThreadPoolExecutor để xử lý các link song song
    max_workers = 5
    with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
        for subject in subjects:
            subject_name = subject["name"]
            subject_url = subject["url"]
            folder_save = subject["folder_save"]

            print("\n======================================")
            print("Đang xử lý subject: {} - URL: {}".format(subject_name, subject_url))
            print("======================================")

            try:
                driver = create_driver()
                driver.get(subject_url)
            except Exception as e:
                logging.error(f"Lỗi khi mở URL {subject_url}: {e}", exc_info=True)
                print(f"  >> Bỏ qua URL {subject_url}")
                continue

            try:
                # Chờ cho đến khi các link với class 'leaf2' xuất hiện
                links = WebDriverWait(driver, 10).until(
                    EC.presence_of_all_elements_located((By.XPATH, "//a[contains(@class, 'leaf2')]"))
                )
            except Exception as e:
                logging.error(f"Lỗi khi chờ link của subject {subject_url}: {e}", exc_info=True)
                print(f"  >> Không tìm thấy link cho subject {subject_url}, bỏ qua.")
                driver.quit()
                continue

            print("Tìm thấy {} link cho subject {}.".format(len(links), subject_name))
            for link in links:
                href = link.get_attribute("href")
                text = link.text.strip()
                if href and text:
                    print("Đang submit xử lý link: {}".format(href))
                    executor.submit(crawl_link, subject_name, folder_save, href, text)
            driver.quit()
    print("\n=== Crawl hoàn thành ===")

if __name__ == '__main__':
    main()
